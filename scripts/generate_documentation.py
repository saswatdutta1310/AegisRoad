"""Generate AegisRoad_Complete_Workflow_Guide.docx — full platform documentation."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AegisRoad_Complete_Workflow_Guide.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_feature(doc, name, purpose, users, steps, notes=None):
    """Structured feature block for consistent documentation."""
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    p.add_run(f" — {purpose}")
    add_para(doc, f"Who uses it: {users}")
    add_para(doc, "How to use it:")
    add_numbered(doc, steps)
    if notes:
        add_para(doc, f"Note: {notes}")


def add_pov_workflow(doc, role_name, steps):
    add_heading(doc, role_name, 2)
    for i, (phase, detail) in enumerate(steps, 1):
        add_heading(doc, f"{i}. {phase}", 3)
        add_para(doc, detail)


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# TITLE PAGE
t = doc.add_heading("AegisRoad", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Complete Workflow, Roles & Feature Guide")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
meta = doc.add_paragraph("Version 3.0  |  Civic Infrastructure Automation Platform\n"
                         "Every role • Every feature • Every workflow perspective")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 1 — PLATFORM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 1 — What Is AegisRoad?", 1)

add_para(doc, (
    "AegisRoad is a closed-loop municipal road-safety and spend-audit platform. It connects "
    "everyone involved in fixing broken roads — from the citizen who spots a pothole, to the "
    "government officer who assigns work, to the contractor who dispatches crews, to the field "
    "worker who pours asphalt and uploads proof. Nothing falls through the cracks: every hazard "
    "has a status, every contractor has an SLA timer, and every repair can be verified with evidence."
))

add_heading(doc, "1.1 The Problem AegisRoad Solves", 2)
add_bullets(doc, [
    "Potholes and cracks often go unreported for months.",
    "Manual assignment to contractors is slow and opaque.",
    "Contractors miss deadlines without financial consequences.",
    "Governments cannot easily prove repairs were done before paying.",
    "Citizens have no visibility into where their tax money goes.",
])

add_heading(doc, "1.2 The Closed-Loop Solution", 2)
add_para(doc, "AegisRoad creates one continuous pipeline:")
add_numbered(doc, [
    "DETECT — Edge AI cameras on buses, citizen apps, or patrol officers log hazards with GPS.",
    "TRIAGE — Government Command Center reviews severity, assigns contractor, starts SLA clock.",
    "DISPATCH — Contractor Portal shows jobs, penalty risk, and fleet allocation tools.",
    "EXECUTE — Field worker navigates to site, logs arrival, uploads before/after photos.",
    "AUDIT — Spend Watch verifies spend, updates contractor scores, releases or withholds payment.",
])

add_heading(doc, "1.3 All Stakeholders (Who Uses the System)", 2)
stakeholders = [
    ("Citizens / Public", "Report hazards via Citizen Report or Landing Page inquiry. Browse public map."),
    ("Edge AI / Dashcam Systems", "Automatically scan road images and POST defects (D00–D40 classes)."),
    ("Government Officers", "Command Center triage, Spend Watch audits, SLA escalations, contractor assignment."),
    ("Contractor Managers", "View assigned jobs, SLA countdown, financial penalty exposure, dispatch crews."),
    ("Field Workers (Drivers / Crew)", "Mobile HUD navigation, arrival logging, photo evidence, mark resolved."),
    ("AegisChat AI Assistant", "Answers questions about hazards, budgets, and compliance for any logged-in user."),
    ("Guests (Observer Mode)", "Browse overview, public map, Edge AI demo, and report issues without login."),
]
for name, desc in stakeholders:
    p = doc.add_paragraph()
    p.add_run(name + ": ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 2 — ROLES IN DEPTH
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 2 — Every Role Explained in Detail", 1)

# --- GOVERNMENT ---
add_heading(doc, "2.1 Government Officer (Municipal / Chief Inspector)", 2)
add_para(doc, (
    "The government role is the master control tower. Officers see every hazard in the city, "
    "decide which contractor fixes what, monitor SLA breaches, and audit whether taxpayer money "
    "was spent correctly. This role represents municipal road departments, chief inspectors, "
    "and infrastructure audit teams."
))

add_heading(doc, "Government — Navigation & Access", 3)
add_bullets(doc, [
    "Login: Use Quick Login 'Chief Inspector Rao' or register with role 'Gov Officer'.",
    "Tabs visible after login: Overview, Command Center, Spend Watch, Hazard Map, Edge AI, Report Issue.",
    "Cannot access: Contractor Portal (contractor-only), Driver Mobile (worker-only).",
])

add_heading(doc, "Government — Daily Responsibilities", 3)
add_numbered(doc, [
    "Monitor live hazard queue in Command Center — filter by severity (critical/high/medium) and status.",
    "Assign unassigned hazards to qualified contractors (e.g. BuildFast Pvt. Ltd.).",
    "Start or track SLA countdown timers — D40 potholes get 24 hours; cracks get up to 96 hours.",
    "Escalate SLA breaches — send Level-4 notifications or re-assign to another firm (e.g. Apex Infrastruct).",
    "Review Spend Watch — check budget allocated vs disbursed per sector and contractor efficiency.",
    "Verify completed jobs have photo evidence before clearing payment.",
    "Use AegisChat to query: 'Any critical unassigned hazards?' or 'BuildFast compliant?'",
])

add_heading(doc, "Government — Key Screens", 3)
gov_screens = {
    "Command Center": "Primary operations desk. KPI cards (active hazards, SLA breaches, contractors online). Hazard table with search/filter. Selected hazard detail panel with Assign Contractor, Mark In Progress, Mark Completed buttons. Live SLA breach list with Escalate and Re-assign actions. Embedded GIS map.",
    "Spend Watch": "Financial ledger. Total allocated/disbursed/remaining budget. Bar charts by sector. Contractor leaderboard ranked by success rate, SLA breaches, penalties. Contract milestone tracking.",
    "Hazard Map (Explorer)": "Public transparency map — all hazards plotted. Citizens and officers can see repair status. Government users can trigger status updates from map popup.",
    "Edge AI": "Review AI-detected defects from uploaded dashcam images before adding to official queue.",
    "Report Issue": "Same as citizen form — officers can manually log hazards from field inspections.",
}
for screen, desc in gov_screens.items():
    p = doc.add_paragraph()
    p.add_run(screen + ": ").bold = True
    p.add_run(desc)

# --- CONTRACTOR ---
add_heading(doc, "2.2 Contractor Manager (Private Construction Firm)", 2)
add_para(doc, (
    "Contractors are private companies hired by the municipality to physically repair roads. "
    "The Contractor Portal is their command hub: they see only jobs assigned to their firm, "
    "watch SLA timers tick down (missing deadlines triggers financial penalties), dispatch "
    "field crews, request materials, and submit photo proof to unlock payment."
))

add_heading(doc, "Contractor — Navigation & Access", 3)
add_bullets(doc, [
    "Login: Quick Login 'Sandra Arjun' (BuildFast Pvt. Ltd.) or register as Contractor.",
    "Only tab after login: Contractor Portal (plus logout).",
    "Jobs filtered by orgName — BuildFast only sees BuildFast assignments.",
    "Guest users see 'Access Denied' on Contractor Portal until they log in.",
])

add_heading(doc, "Contractor — Daily Responsibilities", 3)
add_numbered(doc, [
    "Open Contractor Portal and review Assigned Task Queue.",
    "Monitor Financial Risk Forecaster — total penalty exposure and hourly burn rate.",
    "Select a job and track SLA countdown (simulated live timer per job).",
    "Use Quick Actions: Request SLA Extension (weather), Request Materials (asphalt/cement), Open Support Ticket.",
    "Dispatch crew via Fleet Tracker simulation — see crew availability, distance, ETA.",
    "Update job progress slider (0–100%) as work proceeds.",
    "Upload before/after photo evidence via drag-and-drop zone.",
    "Submit proof to mark job COMPLETED — clears SLA timer for that hazard.",
])

add_heading(doc, "Contractor — Firm Names in System", 3)
add_bullets(doc, [
    "BuildFast Pvt. Ltd. — Primary demo contractor (Sandra Arjun)",
    "Apex Infrastruct — High-score backup contractor (used in re-assign flows)",
    "Core Asphalt Co., Urban Safety Ltd. — Additional selectable firms at registration",
    "BuildRight Co., Z-Force Roads — Appear in hazard data and Spend Watch rankings",
])

# --- FIELD WORKER ---
add_heading(doc, "2.3 Field Worker / Driver (On-Site Crew)", 2)
add_para(doc, (
    "Field workers are the people in trucks with shovels and asphalt. They use a rugged, "
    "high-contrast mobile console designed for tablets mounted in vehicles. They do not "
    "manage budgets or assign jobs — they execute assigned repairs and prove completion."
))

add_heading(doc, "Field Worker — Navigation & Access", 3)
add_bullets(doc, [
    "Login: Quick Login 'Sanjay Kumar' (Eagle Eye Patrols / field fleet).",
    "Only tab after login: Driver Mobile.",
    "Sees jobs where hazard.contractor matches their organization's assigned work (linked via contractor firm on hazard).",
])

add_heading(doc, "Field Worker — Daily Responsibilities", 3)
add_numbered(doc, [
    "View Shift & Safety Advisory — weather and traffic warnings on route.",
    "Select active job from queue — see turn-by-turn HUD (distance, ETA to hazard GPS).",
    "Tap 'Log Arrival' when on-site — timestamps arrival, sets status to in-progress if was unassigned.",
    "Use Voice Log — simulated hands-free dictation for field notes (glove-friendly).",
    "Drop Quick Hazard Pin — report new obstruction at current GPS while on patrol.",
    "Upload photo evidence (drag-and-drop) — before and after repair images.",
    "Mark Resolved / Submit Proof — completes job, sets status to completed, notifies contractor and government.",
])

add_heading(doc, "Field Worker vs Contractor — Key Difference", 3)
add_para(doc, (
    "The Contractor Manager sits in an office and dispatches teams, manages SLA risk, and submits "
    "audit packages. The Field Worker is on the road doing physical labor. Both can upload evidence "
    "and mark jobs complete, but the Contractor Portal adds fleet/material/SLA tools; Driver Mobile "
    "adds navigation HUD and arrival logging."
))

# --- CITIZEN / GUEST ---
add_heading(doc, "2.4 Citizens & Guest Observers", 2)
add_para(doc, (
    "Anyone without login is in Observer Mode. They contribute to the ecosystem by reporting "
    "hazards and viewing public transparency data, but cannot assign contractors or access "
    "financial portals."
))
add_bullets(doc, [
    "Guest tabs: Overview, Public Map, Edge AI, Report Issue.",
    "Citizen Report: issue type, severity, GPS detect, description → creates hazard in system.",
    "Landing Page inquiry form also creates hazards when submitted.",
    "Public Map shows hazard locations and status — builds citizen trust.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 3 — EVERY FEATURE MODULE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 3 — Every Feature Explained", 1)

add_feature(doc, "Overview (Landing Page)",
    "Public-facing home dashboard with KPIs and platform introduction.",
    "Everyone (guests and all roles)",
    [
        "Open app → lands on Overview tab by default.",
        "View live stats: active hazards, contract value (₹ Cr), compliance %, sensor count.",
        "Read solution cards: Road Audits, Route Warnings, SpendWatch, AegisChat.",
        "Submit 'Project Inquiry' form — creates a new unassigned hazard in the queue.",
        "Use navigation cards to jump to Command Center, Map, etc. (role permitting).",
    ],
    "Inquiry form requires name, email, and defect description."
)

add_feature(doc, "Command Center",
    "Government control tower for real-time hazard operations and SLA enforcement.",
    "Government officers (full access); guests cannot open this tab without gov login",
    [
        "View KPI row: critical count, in-progress, completed today, SLA breach alerts.",
        "Use search bar to find hazards by ID, title, or location.",
        "Filter by severity (all/critical/high/medium) and status (unassigned/in-progress/completed).",
        "Click a hazard row → detail panel shows description, contractor, completion %.",
        "Click 'Assign Contractor' → sets contractor name and moves job forward.",
        "Click 'Mark In Progress' or 'Mark Completed' → updates hazard status globally.",
        "Open 'Dispatch Patrol' modal → manually create hazard from officer observation.",
        "SLA Breach section: Escalate (Level-4 email simulation) or Re-assign to Apex Infrastruct.",
        "Embedded map shows hazard pins color-coded by status.",
    ],
    "Live countdown timer simulates active SLA pressure on critical jobs."
)

add_feature(doc, "Spend Watch",
    "Financial audit dashboard linking public budget to contractor performance.",
    "Government officers",
    [
        "View four KPI tiles: Total Allocated, Disbursed, Remaining, Avg Efficiency.",
        "Study bar chart — budget allocated vs disbursed per sector (Metro-01, NH-65, etc.).",
        "Filter contracts by sector or search by contractor/contract ID.",
        "Click contract row → see milestones, quality rating, timeline, tender value.",
        "Review Contractor Leaderboard — sorted by success rate, SLA breaches, penalties.",
        "Click 'Export Audit' to print/save ledger snapshot.",
        "Map panel shows geographic spend distribution.",
    ],
    "Contract data includes AR-7782 (BuildFast), AR-9104 (Urban Grid), etc."
)

add_feature(doc, "Hazard Map (Hazard Explorer)",
    "Interactive Leaflet GIS map for public transparency and geographic triage.",
    "Guests (read-only), Government (can update status from popup)",
    [
        "Pan/zoom dark CartoDB map with hazard markers.",
        "Filter markers: ALL, unassigned, in-progress, completed.",
        "Click pin → popup with title, severity, status, contractor.",
        "Government users: 'Start Repair' button sets in-progress.",
        "Report new hazard button opens inline reporting flow.",
    ],
)

add_feature(doc, "Edge AI Camera Simulator",
    "Upload road images for AI defect detection (YOLOv8) and auto-reporting.",
    "Guests and Government",
    [
        "Drag-and-drop or click to upload road/dashcam image.",
        "Click 'Run Inference' → calls /api/predict (HF Space or demo fallback).",
        "View bounding boxes on detected defects (D00, D10, D20, D40).",
        "Click 'Report to Command Center' → creates hazard with auto GPS and class.",
        "Demo mode shows sample detections if inference server offline.",
    ],
)

add_feature(doc, "Citizen Report",
    "Dedicated public form for reporting road hazards.",
    "Guests and Government",
    [
        "Select issue type: Pothole, Road Crack, Waterlogging, Other.",
        "Set severity: Low, Medium, High, Critical.",
        "Click 'Detect My Location' for GPS coordinates.",
        "Enter description and optional name/email.",
        "Submit → receives reference number, hazard enters municipal queue.",
    ],
    "Pothole maps to class D40; cracks to D20; waterlogging to D10."
)

add_feature(doc, "Contractor Portal",
    "Private firm dashboard for job management, SLA risk, and proof submission.",
    "Contractor managers only (login required)",
    [
        "See firm name and stats: active jobs, success rate, penalty exposure.",
        "Financial Risk card: total penalty ₹ at risk, hourly burn rate.",
        "Quick Actions: SLA Extension, Dispatch Crew, Request Materials, Support Ticket.",
        "Assigned Task Queue — only this contractor's hazards.",
        "Select job → SLA timer, progress slider, evidence upload zone.",
        "Drag files to upload proof → Submit to mark COMPLETED.",
        "Fleet Tracker shows simulated truck positions and ETAs.",
        "Material inventory panel shows asphalt/cement stock levels.",
    ],
)

add_feature(doc, "Driver Mobile (Field Operations Console)",
    "Rugged field-worker UI for navigation, arrival, and repair completion.",
    "Field workers only (login required)",
    [
        "Read safety advisory banner (weather, visibility, traffic).",
        "Pick job from list — shows location, severity, SLA time remaining.",
        "HUD panel: distance to target, ETA, compass-style directions.",
        "Tap 'Log Arrival' — records on-site timestamp, starts work officially.",
        "Tap microphone icon for Voice Log (3-second simulated transcription).",
        "Tap 'Drop Hazard Pin' for quick GPS report while patrolling.",
        "Upload evidence photos via drag-and-drop.",
        "Submit to mark task complete — updates global hazard status.",
    ],
)

add_feature(doc, "AegisChat (Floating AI Assistant)",
    "Claude-powered conversational auditor for hazards, budgets, and SLAs.",
    "All users (bottom-right green chat button)",
    [
        "Click chat bubble to open panel.",
        "Use Quick Queries: NH65 potholes, BuildFast compliance, budget utilisation, SLA alerts.",
        "Type natural language questions about road conditions or contractors.",
        "AI reads live hazard and contract context from the app state.",
        "Requires ANTHROPIC_API_KEY in backend/.env for live responses.",
    ],
)

add_feature(doc, "Authentication System",
    "Role-based login, registration, and demo quick-access accounts.",
    "Everyone",
    [
        "Click 'Login / Sign Up' in header.",
        "Choose role: Contractor, Gov Officer, or Field Worker.",
        "Fill username, email, password, organization.",
        "Or use Quick Login buttons for instant demo access.",
        "Session saved in localStorage; JWT stored when backend connected.",
        "Logout clears session and returns to Observer Mode.",
    ],
)

add_feature(doc, "Interactive Map (Shared Component)",
    "Reusable Leaflet map embedded in Command Center, Spend Watch, Driver Mobile.",
    "Multiple portals",
    [
        "Displays hazard pins with color by status (red=open, amber=in-progress, green=completed).",
        "Synchronized with global HazardContext state.",
        "Click markers for popup details.",
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 4 — WORKFLOWS FROM EVERY POINT OF VIEW
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 4 — Complete Workflows (Every Point of View)", 1)

add_heading(doc, "4.1 Master End-to-End Workflow (All Parties)", 2)
add_numbered(doc, [
    "INGESTION: Hazard enters system via Edge AI, Citizen Report, Landing inquiry, Command Center dispatch, or Field Worker's quick pin.",
    "CLASSIFICATION: System assigns defect class (D00–D40) and severity; SLA hours set automatically (D40=24h, D20=48h, D10=72h, D00=96h).",
    "TRIAGE: Government officer sees hazard as 'unassigned' in Command Center, reviews severity, assigns contractor.",
    "CONTRACTOR ACCEPTANCE: Job appears in Contractor Portal; SLA countdown starts; penalty exposure calculated.",
    "DISPATCH: Contractor uses crew/material tools, assigns field unit (simulated fleet tracker).",
    "FIELD EXECUTION: Worker opens Driver Mobile, navigates to GPS, logs arrival, performs repair.",
    "EVIDENCE: Worker or contractor uploads before/after photos.",
    "RESOLUTION: Status → completed; completion → 100%.",
    "AUDIT: Government reviews in Spend Watch; contractor score updated; payment cleared or penalty applied.",
])

add_heading(doc, "4.2 Workflow — Citizen / Public Point of View", 2)
add_pov_workflow(doc, "Citizen Journey", [
    ("Notice a problem", "Citizen drives over a pothole on NH65 or sees a cracked guardrail."),
    ("Report it", "Opens AegisRoad → Report Issue tab OR fills Landing Page inquiry form."),
    ("Provide details", "Selects Pothole, sets severity High, enables GPS, writes description."),
    ("Confirmation", "Receives reference number (e.g. 482917). Report status = unassigned."),
    ("Wait & observe", "Can open Public Map to see pin appear and watch status change to in-progress then completed."),
    ("Transparency", "Spend Watch (if accessible in demo as guest, primarily gov) shows public money spent on repairs."),
])

add_heading(doc, "4.3 Workflow — Government Officer Point of View", 2)
add_pov_workflow(doc, "Government Daily Flow", [
    ("Morning review", "Login as Chief Inspector Rao → Command Center. Scan KPIs: 8 active hazards, 2 SLA breaches."),
    ("Prioritize", "Filter critical + unassigned. Select HAZ-9821 pothole cluster on NH65."),
    ("Assign", "Assign to BuildFast Pvt. Ltd. Status remains unassigned until contractor acknowledges (or officer marks in-progress)."),
    ("Monitor SLA", "Watch countdown timer. If breach appears in SLA list → Escalate to Level-4 or Re-assign to Apex."),
    ("Mid-day audit", "Open Spend Watch → check BuildFast efficiency (94.5%), Z-Force penalties (2.80 Cr)."),
    ("Verify completion", "When field worker marks completed, officer confirms evidence in hazard description [PROOF UPLOADED]."),
    ("Close loop", "Mark Completed in Command Center if not auto-updated. Ledger updated for disbursement."),
])

add_heading(doc, "4.4 Workflow — Contractor Manager Point of View", 2)
add_pov_workflow(doc, "Contractor Daily Flow", [
    ("Login", "Sandra Arjun logs into BuildFast Pvt. Ltd. account → Contractor Portal only."),
    ("Risk check", "Financial Risk shows ₹2.85L penalty exposure if SLAs missed today."),
    ("Review queue", "3 active jobs assigned to BuildFast. Select pothole on NH65 — 01:45:12 SLA remaining."),
    ("Dispatch", "Open Dispatch Crew modal → assign Heavy Asphalt Team Alpha (14 min ETA)."),
    ("Materials", "Check asphalt stock (142 tons). Request more if below threshold via Materials modal."),
    ("Track field", "Field worker Sanjay (Eagle Eye fleet) navigates via Driver Mobile."),
    ("Receive proof", "Worker uploads photos. Contractor reviews, adjusts progress to 100%, submits audit package."),
    ("Job closed", "Status RESOLVED. SLA timer stops. Success rate maintained for next tender."),
])

add_heading(doc, "4.5 Workflow — Field Worker Point of View", 2)
add_pov_workflow(doc, "Field Worker Shift Flow", [
    ("Start shift", "Sanjay Kumar logs in → Driver Mobile. Reads safety advisory: reduced visibility, moderate traffic."),
    ("Get assignment", "Sees BuildFast jobs in queue. Selects NH65 pothole — HUD shows 450m ahead, 6 min ETA."),
    ("Navigate", "Follows turn-by-turn directions on map panel."),
    ("Arrive", "Taps Log Arrival — toast confirms timestamp. Status → in-progress, 5% complete."),
    ("Work", "Crew pours asphalt. Worker uses Voice Log: 'Pothole filled, 12cm depth, lane 2'."),
    ("Document", "Takes before/after photos → drags into evidence zone."),
    ("Complete", "Submits proof → status completed. Contractor and government see update immediately."),
    ("Patrol", "Optional: Drop Hazard Pin for new crack spotted en route to next job."),
])

add_heading(doc, "4.6 Workflow — Edge AI / Automated Detection Point of View", 2)
add_pov_workflow(doc, "Automated Ingestion Flow", [
    ("Capture", "Bus dashcam or fixed CCTV captures road surface image."),
    ("Upload", "Image sent to Edge AI tab or API /api/predict/ endpoint."),
    ("Inference", "YOLOv8 model returns detections: class D40 pothole, 91% confidence, bounding box coordinates."),
    ("Review", "Operator reviews overlay on image in Edge AI UI."),
    ("Report", "Click Report to Command Center → hazard created with lat/lng, class, severity."),
    ("Join queue", "Hazard enters same pipeline as citizen reports — awaiting government assignment."),
])

add_heading(doc, "4.7 Workflow — SLA Breach & Penalty Point of View", 2)
add_para(doc, (
    "SLA (Service Level Agreement) timers are tied to defect class. When time expires without completion:"
))
add_numbered(doc, [
    "Hazard appears in Command Center SLA Breach list.",
    "Government can Escalate — triggers simulated Level-4 email to Municipal Chief Commissioner.",
    "Government can Re-assign — revokes original contractor, routes to Apex Infrastruct.",
    "Spend Watch shows contractor penalty totals (e.g. Z-Force Roads: 2.80 Cr penalties, 18 breaches).",
    "Contractor Portal Financial Risk card increases penalty exposure in real time.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 5 — HAZARD LIFECYCLE & DATA
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 5 — Hazard Lifecycle & Data Model", 1)

add_heading(doc, "5.1 Status States", 2)
add_bullets(doc, [
    "unassigned (open) — Reported but no contractor assigned yet.",
    "in-progress — Contractor/field crew actively working.",
    "completed (resolved) — Repair done, evidence uploaded, SLA cleared.",
])

add_heading(doc, "5.2 Defect Classification (Pavement Engineering)", 2)
classes = [
    ("D00 — Longitudinal Crack", "Low severity, 96-hour SLA", "Crack running along traffic direction."),
    ("D10 — Transverse Crack", "Medium, 72-hour SLA", "Crack across the lane."),
    ("D20 — Alligator Cracking", "High, 48-hour SLA", "Interconnected crack pattern like reptile skin."),
    ("D40 — Pothole", "Critical, 24-hour SLA", "Structural road cavity — highest priority."),
]
for cls, sla, desc in classes:
    p = doc.add_paragraph()
    p.add_run(cls + " ").bold = True
    p.add_run(f"({sla}) — {desc}")

add_heading(doc, "5.3 Sample Contractors & Performance Data", 2)
add_bullets(doc, [
    "Apex Infrastruct — 99.2% success, 0 SLA breaches (top ranked)",
    "BuildFast Pvt. Ltd. — 94.5% success, 1 breach, 0.05 Cr penalties (demo primary)",
    "BuildRight Co. — 74.1% success, 4 breaches, warning status",
    "Z-Force Roads — 48.5% success, 18 breaches, 2.80 Cr penalties (critical)",
])

add_heading(doc, "5.4 How Data Syncs (Frontend ↔ Backend)", 2)
add_para(doc, (
    "HazardContext and SpendContext hold global state. On app load, the app tries to fetch "
    "from FastAPI (localhost:8000). If backend is offline, rich demo data from data.js is used. "
    "New hazards are optimistically added to UI immediately, then saved to API. Updates to "
    "demo IDs (HAZ-9821) stay local; numeric backend IDs sync via PATCH /api/hazards/{id}."
))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PART 6 — TECHNICAL REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 6 — Installation & Technical Reference", 1)

add_heading(doc, "6.1 Quick Start (Windows)", 2)
add_numbered(doc, [
    "Open PowerShell in project folder d:\\Road_Show",
    "Run .\\start-all.ps1 (starts backend port 8000 + frontend port 3000)",
    "Open http://localhost:3000 in browser",
    "Use Quick Login for any role to test full workflow",
])

add_heading(doc, "6.2 Demo Accounts", 2)
add_bullets(doc, [
    "Chief Inspector Rao — government — Municipal Road Corp",
    "Sandra Arjun — contractor — BuildFast Pvt. Ltd.",
    "Sanjay Kumar — worker — Eagle Eye Patrols",
])

add_heading(doc, "6.3 API Endpoints", 2)
apis = [
    ("GET /", "Health check — API running"),
    ("POST /api/auth/register", "Create user with role and orgName"),
    ("POST /api/auth/login", "Returns JWT access_token + user profile"),
    ("GET /api/hazards/", "List all hazards in database"),
    ("POST /api/hazards/", "Create hazard — auto-sets SLA hours by class"),
    ("PATCH /api/hazards/{id}", "Update status, contractor, completion_percent"),
    ("GET /api/contractors/", "Contractor leaderboard with scores"),
    ("POST /api/chat/", "AegisChat — Claude Sonnet 4 responses"),
    ("POST /api/predict/", "Image upload → YOLO defect detection"),
]
for path, desc in apis:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(path + " — ").bold = True
    p.add_run(desc)

add_heading(doc, "6.4 Environment Variables", 2)
add_bullets(doc, [
    "backend/.env — ANTHROPIC_API_KEY, HF_SPACE_URL, SECRET_KEY",
    "Frontend/.env — VITE_API_URL (optional; default uses Vite proxy)",
])

add_heading(doc, "6.5 Project Structure", 2)
add_para(doc, (
    "Road_Show/Frontend — React UI | Road_Show/backend — FastAPI | "
    "start-all.ps1 — launcher | scripts/generate_documentation.py — regenerates this document"
))

add_heading(doc, "6.6 Troubleshooting", 2)
add_bullets(doc, [
    "Blank contractor portal → Must log in as contractor role.",
    "Chat error → Set ANTHROPIC_API_KEY in backend/.env and restart backend.",
    "Edge AI demo mode → Normal when HF inference server not running.",
    "Toasts not showing → Ensure react-toastify CSS loaded (fixed in main.jsx).",
    "Regenerate this Word file → python scripts/generate_documentation.py",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX — SCENARIO WALKTHROUGHS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Appendix A — Full Scenario: One Pothole from Report to Payment", 1)
add_numbered(doc, [
    "Day 1, 09:00 — Citizen reports pothole on NH65 via Report Issue (D40, Critical). Status: unassigned.",
    "Day 1, 09:15 — Chief Inspector Rao sees it in Command Center, assigns BuildFast Pvt. Ltd.",
    "Day 1, 09:20 — Sandra Arjun (BuildFast) sees job in Contractor Portal. SLA: 24 hours. Penalty clock starts.",
    "Day 1, 10:00 — Sandra dispatches Heavy Asphalt Team Alpha via fleet tool.",
    "Day 1, 10:30 — Sanjay Kumar (field worker) opens Driver Mobile, navigates to NH65, logs arrival.",
    "Day 1, 12:00 — Crew repairs pothole. Sanjay uploads before/after photos, marks completed.",
    "Day 1, 14:00 — Inspector Rao verifies proof in Command Center. Spend Watch shows disbursement eligible.",
    "Day 2 — BuildFast success rate maintained. No penalty. Case closed.",
])

add_heading(doc, "Appendix B — Scenario: SLA Breach & Re-assignment", 1)
add_numbered(doc, [
    "Hazard assigned to Z-Force Roads with 24h SLA.",
    "Timer expires — job not completed.",
    "Appears in SLA Breach list in Command Center.",
    "Inspector clicks Escalate → Level-4 notification logged.",
    "Inspector clicks Re-assign → contractor changed to Apex Infrastruct in description.",
    "Spend Watch shows increased penalties for Z-Force Roads.",
    "Apex completes repair under new assignment.",
])

add_heading(doc, "Appendix C — Role vs Feature Access Matrix", 1)
matrix = [
    "Overview — All users",
    "Command Center — Government only",
    "Spend Watch — Government only",
    "Hazard Map — All users (gov can edit status)",
    "Edge AI — Guests + Government",
    "Report Issue — Guests + Government",
    "Contractor Portal — Contractor only (login required)",
    "Driver Mobile — Field worker only",
    "AegisChat — All users (floating button)",
]
for row in matrix:
    doc.add_paragraph(row, style="List Bullet")

def save_doc():
    try:
        doc.save(OUT)
        return OUT
    except PermissionError:
        alt = ROOT / "AegisRoad_Complete_Workflow_Guide_UPDATED.docx"
        doc.save(alt)
        return alt

path = save_doc()
print(f"Created: {path}")
